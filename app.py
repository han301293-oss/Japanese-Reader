import asyncio
import base64
from datetime import datetime
import io
import json
import os
import re
from concurrent.futures import ThreadPoolExecutor
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
import google.generativeai as genai
import requests
import streamlit as st

# ==============================================================================
# IMPORT THƯ VIỆN AN TOÀN VỚI CƠ CHẾ DỰ PHÒNG
# ==============================================================================
try:
    import pykakasi
    KAKASI_AVAILABLE = True
    _kakasi_instance = pykakasi.kakasi()
except Exception:
    KAKASI_AVAILABLE = False
    _kakasi_instance = None

try:
    import edge_tts
    EDGE_TTS_AVAILABLE = True
except Exception:
    EDGE_TTS_AVAILABLE = False

st.set_page_config(
    page_title="Luyện Đọc & Phân Tích Tiếng Nhật JLPT",
    page_icon="🌸",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# Cấu hình giao diện CSS
st.markdown(
    """
<style>
    header[data-testid="stHeader"] {
        background: linear-gradient(135deg, rgba(255, 240, 245, 0.95), rgba(255, 228, 238, 0.95)) !important;
        border-bottom: 1px solid #ffd1dc !important;
        box-shadow: 0 2px 10px rgba(255, 182, 193, 0.25) !important;
        backdrop-filter: blur(8px) !important;
        z-index: 999990 !important;
    }
    @media (prefers-color-scheme: dark) {
        header[data-testid="stHeader"] {
            background: linear-gradient(135deg, rgba(40, 18, 28, 0.95), rgba(30, 12, 22, 0.95)) !important;
            border-bottom: 1px solid #5a2a3a !important;
            box-shadow: 0 2px 10px rgba(255, 117, 140, 0.2) !important;
        }
    }

    .app-title-fixed {
        position: fixed;
        top: 6px;
        left: 20px;
        z-index: 999999;
        pointer-events: none;
    }
    .app-main-title {
        font-size: 1.35rem;
        font-weight: 800;
        color: #d81b60 !important;
        margin: 0;
        letter-spacing: -0.3px;
    }
    .app-sub-title {
        font-size: 0.78rem;
        color: #666 !important;
        margin: 0;
    }
    @media (prefers-color-scheme: dark) {
        .app-main-title { color: #ff80ab !important; }
        .app-sub-title { color: #bbb !important; }
    }

    .block-container {
        padding-top: 4.2rem !important;
        padding-bottom: 7rem !important;
    }

    .sakura-container {
        position: fixed;
        top: 0; left: 0; width: 100%; height: 100%;
        pointer-events: none;
        z-index: 999;
        overflow: hidden;
    }
    .petal {
        position: absolute;
        background-color: #ffb7c5;
        border-radius: 150% 0 150% 0;
        opacity: 0.7;
        animation: fall 10s linear infinite, sway 3s ease-in-out infinite alternate;
    }
    .petal:nth-child(1) { left: 5%; width: 12px; height: 14px; animation-duration: 9s, 3s; animation-delay: 0s; }
    .petal:nth-child(2) { left: 18%; width: 14px; height: 16px; animation-duration: 11s, 4s; animation-delay: 1.5s; opacity: 0.6; }
    .petal:nth-child(3) { left: 32%; width: 10px; height: 12px; animation-duration: 8s, 2.5s; animation-delay: 3s; }
    .petal:nth-child(4) { left: 45%; width: 15px; height: 17px; animation-duration: 12s, 3.5s; animation-delay: 0.5s; opacity: 0.5; }
    .petal:nth-child(5) { left: 58%; width: 11px; height: 13px; animation-duration: 10s, 3s; animation-delay: 2s; }
    .petal:nth-child(6) { left: 72%; width: 13px; height: 15px; animation-duration: 9.5s, 4s; animation-delay: 4s; opacity: 0.65; }
    .petal:nth-child(7) { left: 85%; width: 12px; height: 14px; animation-duration: 11.5s, 3s; animation-delay: 1s; }
    .petal:nth-child(8) { left: 95%; width: 14px; height: 16px; animation-duration: 8.5s, 2.8s; animation-delay: 2.5s; }

    @keyframes fall {
        0% { top: -20px; transform: rotate(0deg); }
        100% { top: 100vh; transform: rotate(360deg); }
    }
    @keyframes sway {
        0% { transform: translateX(0px) rotate(0deg); }
        100% { transform: translateX(35px) rotate(45deg); }
    }

    ruby { 
        font-size: 1.35rem; 
        line-height: 2.5rem; 
        font-family: 'Hiragino Mincho ProN', 'Yu Mincho', 'Meiryo', serif;
        ruby-align: center;
    }
    rt { 
        font-size: 0.75rem; 
        color: #d81b60; 
        font-weight: 600; 
        user-select: none;
    }
    .plain-jp-text { 
        font-size: 1.25rem; 
        line-height: 2.2rem; 
        font-family: 'Hiragino Mincho ProN', 'Yu Mincho', 'Meiryo', serif; 
    }

    .vi-translation-box {
        background: rgba(255, 243, 224, 0.75);
        border-left: 4px solid #ff9800;
        padding: 10px 16px;
        border-radius: 0 8px 8px 0;
        margin-top: 8px;
        font-size: 1.05rem;
        font-weight: 500;
        color: #d84315;
    }
    @media (prefers-color-scheme: dark) {
        .vi-translation-box {
            background: rgba(60, 40, 20, 0.85);
            border-left: 4px solid #ffb74d;
            color: #ffe0b2;
        }
    }

    .sticky-audio-bar {
        position: fixed;
        bottom: 0;
        left: 0;
        right: 0;
        background: linear-gradient(135deg, rgba(255, 245, 248, 0.98), rgba(255, 235, 242, 0.98));
        border-top: 2px solid #ffccd5;
        padding: 10px 16px;
        box-shadow: 0 -4px 18px rgba(255, 182, 193, 0.35);
        z-index: 99990;
        display: flex;
        flex-wrap: wrap;
        align-items: center;
        justify-content: center;
        gap: 12px;
        backdrop-filter: blur(10px);
    }
    @media (prefers-color-scheme: dark) {
        .sticky-audio-bar {
            background: linear-gradient(135deg, rgba(40, 18, 28, 0.98), rgba(25, 10, 20, 0.98));
            border-top: 2px solid #ff758c;
            box-shadow: 0 -4px 18px rgba(255, 117, 140, 0.25);
        }
    }

    .badge-category {
        display: inline-block;
        padding: 3px 8px;
        border-radius: 6px;
        font-size: 0.8rem;
        font-weight: 700;
        margin-bottom: 6px;
    }
    .badge-author { background-color: #e3f2fd; color: #1565c0; }
    .badge-vocab { background-color: #f3e5f5; color: #7b1fa2; }
    .badge-grammar { background-color: #e8f5e9; color: #2e7d32; }

    /* Nút tuyển dụng mở Tab mới */
    .recruitment-link-card {
        display: flex;
        align-items: center;
        justify-content: space-between;
        margin-top: 14px;
        padding: 14px 20px;
        background: linear-gradient(90deg, #fff2f5 0%, #ffffff 100%);
        border: 1.5px solid #ff80ab;
        border-radius: 10px;
        text-decoration: none !important;
        box-shadow: 0 3px 12px rgba(216, 27, 96, 0.08);
        transition: all 0.25s ease;
    }
    .recruitment-link-card:hover {
        border-color: #d81b60;
        background: linear-gradient(90deg, #ffebee 0%, #fff7f9 100%);
        box-shadow: 0 4px 16px rgba(216, 27, 96, 0.16);
        transform: translateY(-2px);
    }
    .recruitment-badge {
        background: #d81b60;
        color: #ffffff !important;
        font-size: 0.75rem;
        font-weight: 800;
        padding: 3px 8px;
        border-radius: 4px;
        margin-right: 10px;
    }
    .recruitment-title {
        flex: 1;
        font-size: 0.95rem;
        font-weight: 700;
        color: #c2185b !important;
    }
    .recruitment-btn {
        background: #d81b60;
        color: #ffffff !important;
        font-size: 0.82rem;
        font-weight: 700;
        padding: 6px 14px;
        border-radius: 20px;
        display: inline-flex;
        align-items: center;
        gap: 6px;
    }
</style>

<div class="sakura-container">
    <div class="petal"></div><div class="petal"></div><div class="petal"></div>
    <div class="petal"></div><div class="petal"></div><div class="petal"></div>
    <div class="petal"></div><div class="petal"></div>
</div>

<div class="app-title-fixed">
    <div class="app-main-title">🌸 Luyện Đọc & Phân Tích Tiếng Nhật JLPT</div>
    <div class="app-sub-title">Furigana chuẩn ngữ pháp, Giọng Tokyo NHK, Luyện thi Dokkai thông minh</div>
</div>
""",
    unsafe_allow_html=True,
)

api_key = st.secrets.get("GEMINI_API_KEY", None)
sheet_webhook_url = st.secrets.get("GOOGLE_SHEET_WEBHOOK_URL", "")

if not api_key:
    api_key = st.sidebar.text_input(
        "🔑 Nhập Gemini API Key của bạn:", type="password"
    )

if "analysis_data" not in st.session_state:
    st.session_state.analysis_data = None
if "current_user_text" not in st.session_state:
    st.session_state.current_user_text = ""
if "raw_audio_b64" not in st.session_state:
    st.session_state.raw_audio_b64 = None

user_text = st.text_area(
    "📋 Dán bài đọc tiếng Nhật vào đây:",
    height=160,
    placeholder=(
        "例：私たちの意識は、言葉とイメージの網の目をふわふわ漂っているようなものである。それが言葉や文章に定着したとき、「考え」というものになる..."
    ),
)

analyze_btn = st.button(
    "🚀 Phân tích bài đọc", type="primary", use_container_width=True
)

# ==============================================================================
# HÀM XỬ LÝ FURIGANA
# ==============================================================================
def token_to_ruby(orig: str, hira: str) -> str:
    if orig == hira or not orig:
        return orig
    has_kanji = bool(re.search(r"[\u4e00-\u9faf\u3400-\u4dbf]", orig))
    if not has_kanji:
        return orig

    prefix_len = 0
    while (
        prefix_len < len(orig)
        and prefix_len < len(hira)
        and orig[prefix_len] == hira[prefix_len]
    ):
        prefix_len += 1
    prefix = orig[:prefix_len]
    orig_rem = orig[prefix_len:]
    hira_rem = hira[prefix_len:]

    suffix_len = 0
    while (
        suffix_len < len(orig_rem)
        and suffix_len < len(hira_rem)
        and orig_rem[-(suffix_len + 1)] == hira_rem[-(suffix_len + 1)]
    ):
        suffix_len += 1

    if suffix_len > 0:
        kanji_part = orig_rem[:-suffix_len]
        rt_part = hira_rem[:-suffix_len]
        suffix = orig_rem[-suffix_len:]
    else:
        kanji_part = orig_rem
        rt_part = hira_rem
        suffix = ""

    if kanji_part and rt_part:
        return f"{prefix}<ruby>{kanji_part}<rt>{rt_part}</rt></ruby>{suffix}"
    return orig


def convert_to_furigana_html(text: str) -> str:
    if not KAKASI_AVAILABLE or not _kakasi_instance:
        return text
    lines = text.split("\n")
    processed_lines = []
    for line in lines:
        if not line.strip():
            processed_lines.append("")
            continue
        try:
            conv = _kakasi_instance.convert(line)
            line_result = "".join(
                token_to_ruby(item.get("orig", ""), item.get("hira", ""))
                for item in conv
            )
            processed_lines.append(line_result)
        except Exception:
            processed_lines.append(line)
    return "<br>".join(processed_lines)


# ==============================================================================
# HÀM XUẤT FILE WORD (.DOCX)
# ==============================================================================
def generate_docx_file(data_obj):
    doc = Document()
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    topic = data_obj.get("summary", {}).get("topic", "Bài đọc tiếng Nhật")
    level = data_obj.get("summary", {}).get("estimated_jlpt_level", "N/A")
    paragraphs = data_obj.get("paragraphs", [])

    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f"🌸 {topic}")
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(216, 27, 96)

    meta_p = doc.add_paragraph()
    meta_run = meta_p.add_run(
        f"Trình độ ước tính: {level} | Ngày tạo:"
        f" {datetime.now().strftime('%d/%m/%Y')}"
    )
    meta_run.font.size = Pt(10)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("-" * 55)

    for idx, p in enumerate(paragraphs):
        para_title = doc.add_paragraph()
        para_title_run = para_title.add_run(f"Đoạn {idx + 1}:")
        para_title_run.font.size = Pt(11)
        para_title_run.font.bold = True
        para_title_run.font.color.rgb = RGBColor(194, 24, 91)

        jp_p = doc.add_paragraph()
        jp_run = jp_p.add_run(p.get("original_text", ""))
        jp_run.font.size = Pt(12)
        jp_run.font.name = "Meiryo"

        vi_p = doc.add_paragraph()
        vi_run_lbl = vi_p.add_run("🇻🇳 Dịch nghĩa: ")
        vi_run_lbl.font.bold = True
        vi_run = vi_p.add_run(p.get("vietnamese_translation", ""))
        vi_run.font.size = Pt(11)
        vi_run.font.italic = True
        vi_run.font.color.rgb = RGBColor(216, 67, 21)

        doc.add_paragraph()

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


# ==============================================================================
# HÀM BỘ NHỚ ĐỆM & ĐA LUỒNG XỬ LÝ NHANH
# ==============================================================================
def determine_question_rules(text: str):
    char_count = len([c for c in text if not c.isspace()])
    if char_count < 350:
        passage_type = "Đoạn văn ngắn (Tanbun)"
        num_intent, num_vocab, num_grammar = 1, 1, 1
    elif char_count <= 800:
        passage_type = "Đoạn văn trung (Chubun)"
        num_intent, num_vocab, num_grammar = 3, 2, 2
    else:
        passage_type = "Bài báo dài (Choubun)"
        num_intent, num_vocab, num_grammar = 5, 3, 2

    total_questions = num_intent + num_vocab + num_grammar
    return (passage_type, char_count, num_intent, num_vocab, num_grammar, total_questions)


def build_system_prompt(passage_type, num_intent, num_vocab, num_grammar, total_questions):
    return f"""
Bạn là Chuyên gia Ngôn ngữ học tiếng Nhật và Giảng viên Luyện thi JLPT cao cấp (Cấp độ N1).
Nhiệm vụ của bạn là tiếp nhận văn bản tiếng Nhật, phân tích chuyên sâu và trả về kết quả DUY NHẤT dưới dạng JSON hợp lệ (không kèm bất kỳ lời dẫn nào ngoài JSON).

ĐẶC TÍNH BÀI ĐỌC:
- Thể loại: {passage_type}
- Bạn PHẢI tạo CHÍNH XÁC {total_questions} CÂU HỎI TRẮC NGHIỆM JLPT theo tỷ lệ:
  * {num_intent} câu Đọc hiểu ý đồ tác giả / Đại ý bài (category: "Ý đồ tác giả")
  * {num_vocab} câu Từ vựng / Hán tự trong ngữ cảnh bài (category: "Từ vựng / Kanji")
  * {num_grammar} câu Ngữ pháp / Liên từ / Quan hệ logic câu (category: "Ngữ pháp")

CÁC NGUYÊN TẮC QUAN TRỌNG:
1. FURIGANA TRONG PARAGRAPHS:
   - Gắn thẻ `<ruby>Kanji<rt>hiragana</rt></ruby>` chuẩn. 
   - Phần Okurigana để NGOÀI thẻ ruby.
   - Thẻ `<rt>` 100% CHỈ CHỨA CHỮ HIRAGANA THUẦN TÚY (tuyệt đối không chứa chữ Hán hay ký tự khác).
2. DỊCH THUẬT: Dịch thoát ý, tự nhiên, chuẩn xác sang tiếng Việt.
3. TỪ VỰNG: Trích xuất 8 đến 15 từ vựng/cụm từ then chốt (Collocations, quán dụng ngữ, động từ ghép).
4. HÁN TỰ (KANJI): 
   - Trích xuất từ 10 đến 20 Chữ Hán (Kanji) tiêu biểu nhất xuất hiện trong bài đọc kèm Âm Hán Việt, On/Kun và ý nghĩa.
   - Nếu bài đọc quá ngắn hoặc có ít chữ Hán, không bắt buộc đạt mốc 10 chữ mà hãy trích xuất mở rộng tối đa tất cả các chữ Hán có giá trị học tập trong bài.
5. NGỮ PHÁP: Trích xuất 3 đến 6 mẫu ngữ pháp trọng tâm, trích dẫn câu trong bài kèm nghĩa tiếng Việt trong ngoặc `( )`.
6. ĐỀ THI JLPT (CỰC KỲ QUAN TRỌNG):
   - TOÀN BỘ "question_text" và 4 lựa chọn trong "options" (A, B, C, D) PHẢI LÀ 100% TIẾNG NHẬT NGUYÊN BẢN (KHÔNG ĐƯỢC CHỨA BẤT KỲ TỪ TIẾNG VIỆT NÀO).
   - "question_vietnamese": Dịch nghĩa câu hỏi sang tiếng Việt.
   - "correct_explanation": Dịch nghĩa và giải thích lý do tại sao câu trả lời đúng được chọn (bằng tiếng Việt). TUYỆT ĐỐI KHÔNG giải thích hay dịch 3 đáp án sai.

JSON Schema bắt buộc:
{{
  "summary": {{ "estimated_jlpt_level": "N2", "topic": "Chủ đề bài đọc", "word_count": 180 }},
  "paragraphs": [
    {{
      "paragraph_id": 1,
      "original_text": "văn bản gốc",
      "furigana_html": "văn bản có thẻ <ruby>",
      "vietnamese_translation": "bản dịch tiếng Việt"
    }}
  ],
  "grammar_analysis": [
    {{
      "pattern": "mẫu ngữ pháp",
      "jlpt_level": "N3",
      "meaning": "ý nghĩa",
      "usage_in_text": "câu trong bài (nghĩa)",
      "explanation": "giải thích chi tiết"
    }}
  ],
  "vocabulary_list": [
    {{ "word": "từ vựng", "reading": "cách đọc", "jlpt_level": "N2", "vietnamese_meaning": "nghĩa", "part_of_speech": "từ loại" }}
  ],
  "kanji_list": [
    {{ "kanji": "hán tự", "han_viet": "ÂM HÁN", "jlpt_level": "N2", "onyomi": "On", "kunyomi": "Kun", "meaning": "nghĩa" }}
  ],
  "jlpt_practice_questions": [
    {{
      "question_number": 1,
      "category": "Ngữ pháp",
      "question_text": "文章中の「〜」の文法的説明として、最も適切なものはどれか。",
      "question_vietnamese": "Dịch nghĩa câu hỏi bằng tiếng Việt",
      "options": {{
        "A": "「〜」の意味で、原因を表す表現。",
        "B": "「〜」の意味で、逆接を表す表現。",
        "C": "「〜」の意味で、仮定の条件を表す表現。",
        "D": "「〜」の意味で、状態の否定を表す表現。"
      }},
      "correct_answer": "A",
      "correct_explanation": "Dịch nghĩa và giải thích ngắn gọn bằng tiếng Việt vì sao đáp án này đúng."
    }}
  ]
}}
"""


def clean_and_parse_json(raw_text):
    text = raw_text.strip()
    if text.startswith("```json"):
        text = text[7:]
    elif text.startswith("```"):
        text = text[3:]
    if text.endswith("```"):
        text = text[:-3]
    text = text.strip()

    try:
        return json.loads(text, strict=False)
    except Exception:
        cleaned = re.sub(
            r"[\x00-\x1f\x7f-\x9f]",
            lambda m: (" " if m.group(0) not in ["\n", "\r", "\t"] else m.group(0)),
            text,
        )
        try:
            return json.loads(cleaned, strict=False)
        except Exception:
            return None


def clean_text_for_tts(text):
    t = re.sub(r"<[^>]+>", "", text)
    t = re.sub(r"[\r\n]+", " ", t)
    t = re.sub(r"\s+", " ", t)
    t = re.sub(r"[\*\_#`]", "", t)
    return t.strip()


async def _fetch_edge_tts(clean_text):
    voice = "ja-JP-NanamiNeural"
    communicate = edge_tts.Communicate(clean_text, voice)
    mp3_data = bytearray()
    async for chunk in communicate.stream():
        if chunk["type"] == "audio":
            mp3_data.extend(chunk["data"])
    return bytes(mp3_data)


def generate_nhk_voice_sync(text):
    if not EDGE_TTS_AVAILABLE:
        return None
    clean_text = clean_text_for_tts(text)
    if not clean_text:
        return None
    try:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        audio_bytes = loop.run_until_complete(_fetch_edge_tts(clean_text))
        loop.close()
        return audio_bytes
    except Exception:
        return None


# Cache kết quả AI để tăng tốc tức thì khi phân tích lại
@st.cache_data(show_spinner=False, ttl=86400)
def fetch_gemini_analysis(text: str, key: str):
    p_type, _, n_int, n_voc, n_gra, tot_q = determine_question_rules(text)
    genai.configure(api_key=key)
    model = genai.GenerativeModel(
        model_name="gemini-3.6-flash",
        generation_config={
            "response_mime_type": "application/json",
            "temperature": 0.25, # Giảm nhiệt độ để phản hồi nhanh và chuẩn xác hơn
        },
    )
    prompt = build_system_prompt(p_type, n_int, n_voc, n_gra, tot_q)
    response = model.generate_content(f"{prompt}\n\nVăn bản cần phân tích:\n{text}")
    return clean_and_parse_json(response.text)


@st.cache_data(show_spinner=False, ttl=86400)
def fetch_nhk_audio(text: str):
    return generate_nhk_voice_sync(text)


# ==============================================================================
# XỬ LÝ PHÂN TÍCH SONG SONG (PARALLEL EXECUTION)
# ==============================================================================
if analyze_btn:
    if not api_key:
        st.error("⚠️ Vui lòng cấu hình Gemini API Key để tiếp tục.")
    elif not user_text.strip():
        st.warning("⚠️ Vui lòng dán nội dung bài đọc trước khi bấm phân tích.")
    else:
        p_type, char_len, _, _, _, tot_q = determine_question_rules(user_text)

        # Cập nhật thông báo: "Đang phân tích bài văn"
        with st.spinner("⚡ Đang phân tích bài văn..."):
            try:
                # Gọi đồng thời cả AI và tạo Giọng đọc trên 2 luồng riêng biệt
                with ThreadPoolExecutor(max_workers=2) as executor:
                    future_ai = executor.submit(fetch_gemini_analysis, user_text, api_key)
                    future_audio = executor.submit(fetch_nhk_audio, user_text)

                    result = future_ai.result()
                    audio_data = future_audio.result()

                if not result:
                    st.error("⚠️ AI trả về dữ liệu chưa chuẩn cấu trúc. Vui lòng bấm phân tích lại.")
                else:
                    if "paragraphs" in result and isinstance(result["paragraphs"], list):
                        for p in result["paragraphs"]:
                            if isinstance(p, dict):
                                orig = p.get("original_text", "")
                                if KAKASI_AVAILABLE:
                                    p["furigana_html"] = convert_to_furigana_html(orig)

                    st.session_state.analysis_data = result
                    st.session_state.current_user_text = user_text

                if audio_data:
                    st.session_state.raw_audio_b64 = base64.b64encode(audio_data).decode()
                else:
                    st.session_state.raw_audio_b64 = None

            except Exception as e:
                st.error(f"Đã xảy ra lỗi: {str(e)}")

# ==============================================================================
# HIỂN THỊ KẾT QUẢ PHÂN TÍCH (NẾU ĐÃ CÓ DỮ LIỆU)
# ==============================================================================
if st.session_state.analysis_data and isinstance(
    st.session_state.analysis_data, dict
):
    data = st.session_state.analysis_data
    summary_data = (
        data.get("summary", {})
        if isinstance(data.get("summary"), dict)
        else {}
    )
    questions = (
        data.get("jlpt_practice_questions", [])
        if isinstance(data.get("jlpt_practice_questions"), list)
        else []
    )

    st.success("🎉 Đã phân tích thành công!")
    sum_col1, sum_col2, sum_col3 = st.columns(3)
    sum_col1.info(
        f"🏷️ **Cấp độ ước tính:** {summary_data.get('estimated_jlpt_level', 'N/A')}"
    )
    sum_col2.info(f"📖 **Chủ đề:** {summary_data.get('topic', 'Chung')}")
    sum_col3.info(f"📊 **Đề thi:** {len(questions)} câu hỏi JLPT")

    tab_read, tab_grammar, tab_vocab, tab_kanji, tab_quiz = st.tabs([
        "📖 Bài đọc & Dịch",
        "📝 Ngữ pháp trọng tâm",
        "📚 Từ vựng then chốt",
        "🈲 Hán tự (Kanji)",
        f"❓ Đề thi JLPT ({len(questions)} câu)",
    ])

    with tab_read:
        ctrl_col1, ctrl_col2 = st.columns([1.2, 3.8])
        with ctrl_col1:
            # Mặc định tắt Furigana, người dùng cần sẽ bật lên sau
            show_furigana = st.toggle("🌸 Bật Furigana", value=False)
        with ctrl_col2:
            docx_file = generate_docx_file(data)
            topic_slug = re.sub(
                r"[^\w\s-]", "", summary_data.get("topic", "BaiDoc")
            ).strip()
            topic_slug = re.sub(r"[-\s]+", "_", topic_slug)
            st.download_button(
                label="📥 Tải bài đọc & bản dịch (.docx / Word)",
                data=docx_file,
                file_name=f"JLPT_Reading_{topic_slug}.docx",
                mime=(
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
                type="secondary",
            )

        st.markdown("---")

        paragraphs = (
            data.get("paragraphs", [])
            if isinstance(data.get("paragraphs"), list)
            else []
        )
        for p in paragraphs:
            if isinstance(p, dict):
                if show_furigana:
                    st.markdown(
                        f"<div>{p.get('furigana_html', '')}</div>",
                        unsafe_allow_html=True,
                    )
                else:
                    st.markdown(
                        f"<div class='plain-jp-text'>{p.get('original_text', '')}</div>",
                        unsafe_allow_html=True,
                    )
                st.markdown(
                    "<div class='vi-translation-box'>🇻🇳 <strong>Dịch"
                    f" nghĩa:</strong> {p.get('vietnamese_translation', '')}</div>",
                    unsafe_allow_html=True,
                )
                st.markdown(
                    "<div style='margin-bottom: 15px;'></div>",
                    unsafe_allow_html=True,
                )

        if st.session_state.raw_audio_b64:
            st.markdown(
                f"""
            <div class="sticky-audio-bar">
                <span style="font-size: 0.9rem; font-weight: 700; color: #d81b60;">🎙️ Giọng đọc chuẩn Tokyo (NHK):</span>
                <div style="display: flex; align-items: center; gap: 6px;">
                    <button onclick="let p = document.getElementById('floating_player'); if(p) p.currentTime = Math.max(0, p.currentTime - 10);" 
                            style="padding: 4px 10px; border-radius: 8px; border: 1px solid #ffccd5; background: white; font-weight: 700; font-size: 0.82rem; cursor: pointer; color: #d81b60;">
                        ⏮️ -10s
                    </button>
                    <button onclick="let p = document.getElementById('floating_player'); if(p) p.currentTime = Math.min(p.duration || 0, p.currentTime + 10);" 
                            style="padding: 4px 10px; border-radius: 8px; border: 1px solid #ffccd5; background: white; font-weight: 700; font-size: 0.82rem; cursor: pointer; color: #d81b60;">
                        +10s ⏭️
                    </button>
                </div>
                <audio id="floating_player" controls style="height: 38px; max-width: 360px; flex-grow: 1;">
                    <source src="data:audio/mp3;base64,{st.session_state.raw_audio_b64}" type="audio/mp3">
                </audio>
                <div style="display: flex; align-items: center; gap: 6px;">
                    <span style="font-size: 0.82rem; font-weight: 600;">⚡ Tốc độ:</span>
                    <select id="speed_select" onchange="document.getElementById('floating_player').playbackRate = this.value;" style="padding: 4px 8px; border-radius: 8px; border: 1px solid #ffccd5; background: white; font-weight: 600; cursor: pointer;">
                        <option value="0.5">x0.5 (Rất chậm)</option>
                        <option value="0.75">x0.75 (Chậm)</option>
                        <option value="1.0" selected>x1.0 (Chuẩn)</option>
                        <option value="1.25">x1.25 (Nhanh vừa)</option>
                        <option value="1.5">x1.5 (Nhanh)</option>
                        <option value="1.75">x1.75 (Rất nhanh)</option>
                        <option value="2.0">x2.0 (Cực nhanh)</option>
                    </select>
                </div>
            </div>
            """,
                unsafe_allow_html=True,
            )

    with tab_grammar:
        grammars = (
            data.get("grammar_analysis", [])
            if isinstance(data.get("grammar_analysis"), list)
            else []
        )
        if not grammars:
            st.info("Không có mẫu ngữ pháp đặc biệt nào.")
        for g in grammars:
            if isinstance(g, dict):
                p_text = g.get("pattern", "")
                l_text = g.get("jlpt_level", "")
                m_text = g.get("meaning", "")
                exp_title = f"📌 {p_text} [{l_text}] — {m_text}"
                with st.expander(exp_title, expanded=True):
                    st.markdown(
                        f"- **Ngữ cảnh trong bài:** `{g.get('usage_in_text', '')}`"
                    )
                    st.markdown(
                        f"- **Giải thích chi tiết:** {g.get('explanation', '')}"
                    )

    with tab_vocab:
        vocabs = (
            data.get("vocabulary_list", [])
            if isinstance(data.get("vocabulary_list"), list)
            else []
        )
        if vocabs:
            # Chuyển Cấp độ và Từ loại ra phía sau cột Ý nghĩa trong bài
            vocab_rows = [
                {
                    "Từ vựng / Cụm từ": v.get("word", ""),
                    "Cách đọc (Kana)": v.get("reading", ""),
                    "Ý nghĩa trong bài": v.get("vietnamese_meaning", ""),
                    "Cấp độ": v.get("jlpt_level", ""),
                    "Từ loại": v.get("part_of_speech", ""),
                }
                for v in vocabs
                if isinstance(v, dict)
            ]
            st.dataframe(vocab_rows, use_container_width=True, hide_index=True)
        else:
            st.info("Chưa có danh sách từ vựng.")

    with tab_kanji:
        kanjis = (
            data.get("kanji_list", [])
            if isinstance(data.get("kanji_list"), list)
            else []
        )
        if kanjis:
            # Chuyển Cấp độ, Âm On, Âm Kun ra phía sau cột Ý nghĩa
            kanji_rows = [
                {
                    "Hán tự": k.get("kanji", ""),
                    "Âm Hán Việt": k.get("han_viet", ""),
                    "Ý nghĩa": k.get("meaning", ""),
                    "Cấp độ": k.get("jlpt_level", ""),
                    "Âm On": k.get("onyomi", ""),
                    "Âm Kun": k.get("kunyomi", ""),
                }
                for k in kanjis
                if isinstance(k, dict)
            ]
            st.dataframe(kanji_rows, use_container_width=True, hide_index=True)
        else:
            st.info("Chưa có danh sách Hán tự.")

    with tab_quiz:
        st.markdown(f"### ✍️ Đề thi thử JLPT ({len(questions)} câu hỏi)")
        for idx, q in enumerate(questions):
            if isinstance(q, dict):
                q_num = q.get("question_number", idx + 1)
                category = q.get("category", "Đọc hiểu")

                badge_class = "badge-author"
                if "từ vựng" in category.lower() or "kanji" in category.lower():
                    badge_class = "badge-vocab"
                elif "ngữ pháp" in category.lower():
                    badge_class = "badge-grammar"

                st.markdown(
                    f"<span class='badge-category {badge_class}'>🏷️"
                    f" {category}</span>",
                    unsafe_allow_html=True,
                )
                st.markdown(
                    f"#### Câu {q_num}: {q.get('question_text', '')}"
                )

                opts = (
                    q.get("options", {})
                    if isinstance(q.get("options"), dict)
                    else {}
                )
                choice_keys = [k for k in ["A", "B", "C", "D"] if k in opts]

                user_choice = st.radio(
                    f"Chọn phương án đúng cho câu {q_num}:",
                    options=choice_keys,
                    format_func=lambda x: f"{x}. {opts.get(x, '')}",
                    key=f"quiz_radio_{q_num}",
                    index=None,
                )

                correct_ans = q.get("correct_answer", "A")

                # Chỉ khi người dùng chọn đáp án mới hiển thị bản dịch câu hỏi và giải thích câu đúng
                if user_choice is not None:
                    st.caption(f"*(Dịch nghĩa câu hỏi: {q.get('question_vietnamese', '')})*")

                    if user_choice == correct_ans:
                        st.success(
                            "🎉 **Chính xác!** Đáp án đúng là"
                            f" **{correct_ans}**."
                        )
                    else:
                        st.error(
                            f"❌ **Chưa chính xác!** Bạn chọn **{user_choice}**,"
                            f" đáp án chuẩn là **{correct_ans}**."
                        )

                    correct_opt_text = opts.get(correct_ans, "")
                    explanation = q.get("correct_explanation") or q.get("option_analysis", {}).get(correct_ans, "Đáp án chuẩn xác.")
                    
                    st.markdown("**🔍 Dịch nghĩa & Giải thích đáp án đúng:**")
                    st.info(f"✅ **Đáp án {correct_ans}:** {correct_opt_text}\n\n👉 **Ý nghĩa & Giải thích:** {explanation}")
                st.markdown("---")

# ==============================================================================
# MỤC 1: FORM PHẢN HỒI (LUÔN HIỂN THỊ Ở CHÂN TRANG)
# ==============================================================================
st.markdown("<br><hr>", unsafe_allow_html=True)
with st.expander("💌 Góp ý & Báo lỗi"):
    with st.form("feedback_form", clear_on_submit=True):
        fb_name = st.text_input("Tên hoặc Email (không bắt buộc):")
        fb_type = st.selectbox(
            "Loại góp ý:",
            [
                "Báo lỗi Furigana / AI",
                "Lỗi giọng đọc",
                "Đề xuất tính năng mới",
                "Khác",
            ],
        )
        fb_content = st.text_area(
            "Nội dung:", placeholder="Mô tả ý kiến của bạn..."
        )
        submitted = st.form_submit_button("📩 Gửi ý kiến")
        if submitted and fb_content.strip():
            if sheet_webhook_url:
                try:
                    payload = {
                        "time": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
                        "name": fb_name.strip() if fb_name.strip() else "Ẩn danh",
                        "type": fb_type,
                        "content": fb_content,
                    }
                    requests.post(sheet_webhook_url, json=payload, timeout=5)
                except Exception:
                    pass
            st.success(
                "🌸 Cảm ơn bạn! Ý kiến đóng góp đã được gửi thành công."
            )

# ==============================================================================
# MỤC 2: NÚT TUYỂN DỤNG ĐIỀU HƯỚNG TỚI TRANG CON CỦA STREAMLIT
# ==============================================================================
st.markdown(
    """
    <a href="/Tuyển_Dụng" target="_blank" rel="noopener noreferrer" class="recruitment-link-card">
        <div style="display: flex; align-items: center;">
            <span class="recruitment-badge">HOT</span>
            <span class="recruitment-title">🔥 TUYỂN DỤNG NHÂN SỰ TIẾNG NHẬT TỪ N3 — KHÔNG YÊU CẦU KINH NGHIỆM</span>
        </div>
        <div class="recruitment-btn">
            Xem chi tiết ➔
        </div>
    </a>
    """,
    unsafe_allow_html=True,
)
